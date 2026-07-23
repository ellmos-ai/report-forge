#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.

document_extraction.py -- Domaenen-neutrale Quelldokument-Extraktion
=======================================================================

Portiert aus dem privaten foerderplaner-Skill (generator/generator.py,
Phase-1-Extraktion). Liest Quelldokumente (Word, PDF, Text, Excel, Mail)
in reinen Text um -- ohne jede domaenenspezifische Annahme ueber Inhalt
oder Struktur. Wird von ReportWorkflow.prepare() vor der Anonymisierung/
Prompt-Erstellung aufgerufen.

Version: 1.0.0
"""

from __future__ import annotations

from pathlib import Path

MAX_SOURCE_FILE_BYTES = 20 * 1024 * 1024
MAX_SOURCE_TEXT_CHARS = 500_000

SUPPORTED_SUFFIXES = {".docx", ".txt", ".md", ".pdf", ".xlsx", ".xls", ".msg", ".eml"}


def extract_text_from_docx(filepath: str) -> str:
    """Extrahiert Text aus einem Word-Dokument (Absaetze + Tabellenzellen)."""
    try:
        from docx import Document
    except ImportError:
        return "[FEHLER: python-docx nicht installiert]"

    doc = Document(filepath)
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_txt(filepath: str) -> str:
    """Liest eine Textdatei (UTF-8 mit Latin-1-Fallback)."""
    path = Path(filepath)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def extract_text_from_pdf(filepath: str) -> str:
    """Extrahiert Text aus einem PDF. pypdf (MIT) primaer, pdfplumber/PyMuPDF optional."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    try:
        import pdfplumber

        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    try:
        import fitz  # PyMuPDF, AGPL -- nur wenn der Nutzer es explizit installiert

        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    except ImportError:
        return "[PDF-Extraktion nicht möglich: kein unterstützter PDF-Reader installiert]"


def extract_text_from_excel(filepath: str) -> str:
    """Extrahiert Text aus einer Excel-Datei (.xlsx, .xls)."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(filepath, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            parts.append(f"[Tabelle: {sheet_name}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[Excel-Extraktion nicht möglich: openpyxl nicht installiert]"
    except Exception as exc:
        return f"[Excel-Fehler ({type(exc).__name__})]"


def extract_text_from_msg(filepath: str) -> str:
    """Extrahiert Text aus einer Outlook .msg Datei."""
    try:
        import extract_msg

        msg = extract_msg.Message(filepath)
        parts = []
        if msg.date:
            parts.append(f"Datum: {msg.date}")
        if msg.sender:
            parts.append(f"Von: {msg.sender}")
        if msg.to:
            parts.append(f"An: {msg.to}")
        if msg.subject:
            parts.append(f"Betreff: {msg.subject}")
        parts.append("")
        if msg.body:
            parts.append(msg.body)
        msg.close()
        return "\n".join(parts)
    except ImportError:
        return "[MSG-Extraktion nicht möglich: extract-msg nicht installiert]"
    except Exception as exc:
        return f"[MSG-Fehler ({type(exc).__name__})]"


def extract_text_from_eml(filepath: str) -> str:
    """Extrahiert Text aus einer .eml Datei."""
    try:
        from email import policy
        from email.parser import BytesParser

        with open(filepath, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []
        if msg["Date"]:
            parts.append(f"Datum: {msg['Date']}")
        if msg["From"]:
            parts.append(f"Von: {msg['From']}")
        if msg["To"]:
            parts.append(f"An: {msg['To']}")
        if msg["Subject"]:
            parts.append(f"Betreff: {msg['Subject']}")
        parts.append("")

        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    charset = part.get_content_charset() or "utf-8"
                    try:
                        body = part.get_payload(decode=True).decode(charset)
                        parts.append(body)
                        break
                    except Exception:
                        pass
        else:
            charset = msg.get_content_charset() or "utf-8"
            try:
                body = msg.get_payload(decode=True).decode(charset)
                parts.append(body)
            except Exception:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)
    except Exception as exc:
        return f"[EML-Fehler ({type(exc).__name__})]"


def extract_all_sources(source_folder: str, recursive: bool = True) -> str:
    """
    Liest alle unterstuetzten Quelldokumente aus einem Ordner und fasst sie
    zu einem einzigen Text zusammen (mit Quellenmarkern je Datei).

    Unterstuetzt: .docx, .txt, .md, .pdf, .xlsx, .xls, .msg, .eml
    Ueberspringt versteckte/Underscore-Dateien und typische Export-/
    Deexport-Verzeichnisse (output, _ready, _archive, _prepare).

    Args:
        source_folder: Pfad zum Quellordner
        recursive: Auch Unterordner durchsuchen (default: True)

    Raises:
        ValueError: Groessenlimit (Einzeldatei oder Gesamttext) ueberschritten
    """
    folder = Path(source_folder).expanduser().resolve()
    if not folder.is_dir():
        return "[FEHLER: Quellordner nicht gefunden]"

    all_text = []
    files = sorted(folder.rglob("*")) if recursive else sorted(folder.iterdir())

    total_chars = 0
    excluded_prefixes = ("output", "_ready", "_archive", "_prepare")
    for filepath in files:
        if filepath.is_symlink():
            continue
        if filepath.name.startswith(".") or filepath.name.startswith("_"):
            continue
        rel_path = filepath.relative_to(folder)
        if any(part.casefold().startswith(excluded_prefixes) for part in rel_path.parts):
            continue
        if not filepath.is_file():
            continue
        suffix = filepath.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue
        if filepath.stat().st_size > MAX_SOURCE_FILE_BYTES:
            raise ValueError("Quelldatei überschreitet das Größenlimit")

        if suffix == ".docx":
            text = extract_text_from_docx(str(filepath))
        elif suffix in (".txt", ".md"):
            text = extract_text_from_txt(str(filepath))
        elif suffix == ".pdf":
            text = extract_text_from_pdf(str(filepath))
        elif suffix in (".xlsx", ".xls"):
            text = extract_text_from_excel(str(filepath))
        elif suffix == ".msg":
            text = extract_text_from_msg(str(filepath))
        elif suffix == ".eml":
            text = extract_text_from_eml(str(filepath))
        else:
            continue

        total_chars += len(text)
        if total_chars > MAX_SOURCE_TEXT_CHARS:
            raise ValueError("Extrahierter Quelltext überschreitet das Größenlimit")
        all_text.append(f"--- Quelle {len(all_text) + 1} ({suffix}) ---\n{text}")

    if not all_text:
        return "[FEHLER: Keine Quelldokumente gefunden]"

    return "\n\n".join(all_text)
