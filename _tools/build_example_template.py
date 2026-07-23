# -*- coding: utf-8 -*-
"""Einmal-Skript: baut templates/example_template.docx (neutral, keine Bilder)."""
from pathlib import Path
from docx import Document

doc = Document()
doc.add_heading("Generischer Bericht (report-forge Beispiel)", level=1)

doc.add_paragraph("Name: {{SUBJECT_NAME}}")
doc.add_paragraph("Referenzdatum: {{SUBJECT_REFERENCE_DATE}}")

doc.add_heading("Befunde", level=2)
table = doc.add_table(rows=2, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Code"
hdr[1].text = "Aussage"
hdr[2].text = "Status"
tpl = table.rows[1].cells
tpl[0].text = ""
tpl[1].text = ""
tpl[2].text = ""

doc.add_heading("Zusammenfassung", level=2)
doc.add_paragraph("{{SUMMARY}}")

doc.add_heading("Empfehlung", level=2)
doc.add_paragraph("Fortsetzung empfohlen: {{RECOMMENDATION_CONTINUE}}")
doc.add_paragraph("Anmerkung: {{RECOMMENDATION_NOTES}}")

out = Path(r"C:\Users\User\OneDrive\.TOPICS\.AI\.MODULES\.DOMAINS\report-forge\templates\example_template.docx")
doc.save(str(out))
print("saved:", out, out.stat().st_size, "bytes")
