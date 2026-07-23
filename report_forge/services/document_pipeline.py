#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.

DocumentPipeline — Domaenen-neutraler Dokumenten-Verarbeitungs-Service
=========================================================================

Portiert und GENERALISIERT aus dem privaten foerderplaner-Skill
(services/document_pipeline.py, v1.2.0). Dort war die Dokumenttyp-
Erkennung hart auf die ICF-/Autismus-Foerderplanung zugeschnitten
(z.B. "proautismus_bericht", "arztbericht", "schulbericht"). Hier ist
die Typ-Erkennung ueber eine austauschbare Muster-Registry konfigurierbar;
die Default-Registry (DEFAULT_DOC_PATTERNS) nutzt neutrale, domaenen-
uebergreifende Kategorien (Protokoll/Log, Aufnahmeblatt, Genehmigung/
Bewilligung, Fachbericht, externer Fachbericht, institutioneller Bericht,
Mail).

Bietet drei Interface-Schichten:
  1. Scan-Interface: scan_folder() -> ScanResult
  2. Extraktions-Interface: extract_bundle() -> TextBundle
  3. Bundle-Lese-Interface: read_bundle_dir() -> str

HINWEIS: Dieser Service ist aktuell NICHT in ReportWorkflow verdrahtet
(der Standard-Extraktionspfad ist document_extraction.extract_all_sources,
identisch zum bisherigen Verhalten des foerderplaner-Skills). Er steht als
eigenstaendiger, fortgeschrittener Baustein fuer Overlays zur Verfuegung,
die Prioritaets-/Kategorisierungslogik (CORE/STUFE2/EXTENDED/SKIP) fuer
grosse, gemischte Dokumentbestaende brauchen.

Version: 1.0.0 (generalisiert aus foerderplaner document_pipeline.py 1.2.0)
"""

import os
import re
import tempfile
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TYPE_CHECKING
from enum import Enum

if TYPE_CHECKING:
    pass


# ═══════════════════════════════════════════════════════════════
# Datenklassen
# ═══════════════════════════════════════════════════════════════


class DocumentCategory(Enum):
    """Dokument-Kategorien nach Prioritaet (domaenen-neutral)."""

    CORE = "core"  # Immer an das LLM senden
    TIER2 = "tier2"  # Nach CORE senden
    EXTENDED = "extended"  # Nur auf Anfrage
    SKIP = "skip"  # Nicht verarbeiten


CATEGORY_PRIORITY = {
    DocumentCategory.CORE: 0,
    DocumentCategory.TIER2: 1,
    DocumentCategory.EXTENDED: 2,
    DocumentCategory.SKIP: 3,
}

# Generische Default-Dokumenttyp-Prioritaet innerhalb derselben Kategorie.
# Overlays koennen eigene Registries injizieren (siehe DocumentPipeline.__init__).
DEFAULT_DOC_TYPE_PRIORITY = {
    "log": 0,
    "intake_sheet": 1,
    "approval": 2,
    "domain_report": 3,
    "professional_report": 4,
    "institutional_report": 5,
    "mail": 6,
    "root_document": 0,
    "unknown": 8,
}


@dataclass
class DocumentInfo:
    """Informationen zu einem Dokument."""

    path: Path
    filename: str
    suffix: str
    category: DocumentCategory
    doc_type: str  # z.B. "log", "approval", "professional_report"
    date_hint: Optional[datetime] = None
    size_bytes: int = 0
    text_extracted: bool = False
    text_length: int = 0


@dataclass
class ScanResult:
    """Ergebnis der Dokumenten-Sammlung."""

    documents: List[DocumentInfo] = field(default_factory=list)
    core_count: int = 0
    tier2_count: int = 0
    extended_count: int = 0
    skipped_count: int = 0
    errors: List[str] = field(default_factory=list)


@dataclass
class TextBundle:
    """Extrahiertes Text-Bundle."""

    core_text: str = ""
    tier2_text: str = ""
    extended_text: str = ""
    total_length: int = 0
    files_processed: int = 0


# ═══════════════════════════════════════════════════════════════
# Default-Erkennungsmuster (generisch, ueberschreibbar)
# ═══════════════════════════════════════════════════════════════

DEFAULT_DOC_PATTERNS: Dict[str, List[str]] = {
    # CORE-Kandidaten
    "log": [
        r"protokoll",
        r"log",
        r"doku(?:mentation)?",
        r"verlauf",
        r"sitzung",
        r"session.?notes?",
    ],
    "intake_sheet": [
        r"aktendeckblatt",
        r"stammdaten",
        r"anmeld(?:ung|ebogen|eunterlagen)",
        r"intake",
    ],
    "approval": [
        r"hilfeplan",
        r"kostenzusage",
        r"bewilligung",
        r"bescheid",
        r"approval",
        r"authoriz",
    ],
    "domain_report": [
        r"entwicklungsbericht",
        r"foerderbericht",
        r"bericht",
        r"report",
    ],
    # TIER2-Kandidaten
    "professional_report": [
        r"arztbericht",
        r"bericht.*dr\.?",
        r"entlass(?:ungs)?bericht",
        r"diagnose",
        r"befund",
        r"gutachten",
        r"assessment",
    ],
    "institutional_report": [
        r"schulbericht",
        r"zeugnis",
        r"lernentwicklung",
        r"nachteilsausgleich",
        r"institutional.?report",
    ],
    "mail": [
        r"\.msg$",
        r"\.eml$",
    ],
}

DEFAULT_FOLDER_HINTS: Dict[str, List[str]] = {
    "core": ["dokumentation", "protokolle", "intake", "approval", "intern"],
    "tier2": ["extern", "professional", "institution", "mail"],
    "extended": ["archiv", "alt", "historie", "backup"],
    "skip": ["output", "_prepare", "_ready", "_archive", "temp", "_bundle"],
}

SUPPORTED_EXTENSIONS = {
    ".docx", ".doc", ".txt", ".md", ".pdf", ".xlsx", ".xls", ".msg", ".eml",
}
BUNDLE_READ_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".xlsx", ".xls"}
MAX_DOCUMENT_FILES = 500
MAX_DOCUMENT_BYTES = 20 * 1024 * 1024
MAX_TOTAL_DOCUMENT_BYTES = 100 * 1024 * 1024


# ═══════════════════════════════════════════════════════════════
# DocumentPipeline — Haupt-Klasse
# ═══════════════════════════════════════════════════════════════


class DocumentPipeline:
    """
    Vereinter, domaenen-neutraler Dokumenten-Verarbeitungs-Service.

    Verwendung (Default-Registry):
        pipeline = DocumentPipeline()
        result = pipeline.scan_folder("/path/to/source")
        bundle = pipeline.extract_bundle(result.documents)

    Verwendung (eigene Domaenen-Registry, z.B. im Overlay):
        pipeline = DocumentPipeline(
            doc_patterns={"clinical_note": [r"klinische.?notiz"], ...},
            doc_type_priority={"clinical_note": 0, "unknown": 8},
        )

    Sicherheitshinweis: Originaldateien werden nie veraendert. Wird ein
    anonym_profile-Objekt mit einer `mappings`-Struktur (Kategorie ->
    {original: ersatz}) uebergeben, wird der extrahierte Text anhand
    dieser Mappings ersetzt -- die konkrete Anonymisierungslogik liegt
    beim aufrufenden Overlay (z.B. dem anonymizer-Modul).
    """

    def __init__(
        self,
        report_period_months: int = 12,
        reference_date: Optional[datetime] = None,
        doc_patterns: Optional[Dict[str, List[str]]] = None,
        doc_type_priority: Optional[Dict[str, int]] = None,
        folder_hints: Optional[Dict[str, List[str]]] = None,
    ):
        """
        Args:
            report_period_months: Zeitraum fuer "aktuelle" Dokumente (Default: 12)
            reference_date: Referenzdatum (Default: heute)
            doc_patterns: Ueberschreibt DEFAULT_DOC_PATTERNS
            doc_type_priority: Ueberschreibt DEFAULT_DOC_TYPE_PRIORITY
            folder_hints: Ueberschreibt DEFAULT_FOLDER_HINTS
        """
        self.report_period_months = report_period_months
        self.reference_date = reference_date or datetime.now()
        self.period_start = self.reference_date - timedelta(
            days=report_period_months * 30
        )
        self.doc_patterns = doc_patterns or DEFAULT_DOC_PATTERNS
        self.doc_type_priority = doc_type_priority or DEFAULT_DOC_TYPE_PRIORITY
        self.folder_hints = folder_hints or DEFAULT_FOLDER_HINTS

        self._compiled_patterns: Dict[str, List[re.Pattern]] = {}
        for doc_type, patterns in self.doc_patterns.items():
            self._compiled_patterns[doc_type] = [
                re.compile(p, re.IGNORECASE) for p in patterns
            ]

    # ─────────────────────────────────────────────────────────────
    # Interface 1: Ordner scannen und kategorisieren
    # ─────────────────────────────────────────────────────────────

    def scan_folder(self, folder: str, recursive: bool = True) -> ScanResult:
        """Scannt einen Ordner und kategorisiert alle Dokumente."""
        result = ScanResult()
        folder_path = Path(folder).expanduser().resolve()

        if not folder_path.is_dir():
            result.errors.append("Quellordner nicht gefunden")
            return result

        files = list(folder_path.rglob("*")) if recursive else list(folder_path.iterdir())

        total_bytes = 0
        accepted_files = 0
        for filepath in sorted(files):
            if filepath.is_symlink() or not filepath.is_file():
                continue
            if filepath.name.startswith("."):
                continue
            if filepath.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            size = filepath.stat().st_size
            if size > MAX_DOCUMENT_BYTES:
                result.errors.append("Quelldatei überschreitet das Größenlimit")
                result.documents.clear()
                return result
            total_bytes += size
            accepted_files += 1
            if accepted_files > MAX_DOCUMENT_FILES or total_bytes > MAX_TOTAL_DOCUMENT_BYTES:
                result.errors.append("Quellordner überschreitet das Ressourcenlimit")
                result.documents.clear()
                return result

            doc_info = self._analyze_document(filepath, folder_path)
            result.documents.append(doc_info)

            if doc_info.category == DocumentCategory.CORE:
                result.core_count += 1
            elif doc_info.category == DocumentCategory.TIER2:
                result.tier2_count += 1
            elif doc_info.category == DocumentCategory.EXTENDED:
                result.extended_count += 1
            else:
                result.skipped_count += 1

        ten_years_ago = self.reference_date - timedelta(days=10 * 365)

        def _sort_key(d: DocumentInfo):
            cat_prio = CATEGORY_PRIORITY.get(d.category, 9)
            type_prio = self.doc_type_priority.get(d.doc_type, 8)
            if d.doc_type == "professional_report" and d.date_hint and d.date_hint < ten_years_ago:
                type_prio = 99
            date_val = -d.date_hint.timestamp() if d.date_hint else 0
            return (cat_prio, type_prio, date_val)

        result.documents.sort(key=_sort_key)
        return result

    def _analyze_document(self, filepath: Path, base_folder: Path) -> DocumentInfo:
        """Analysiert ein einzelnes Dokument und bestimmt Kategorie."""
        filename = filepath.name
        suffix = filepath.suffix.lower()

        try:
            rel_path = filepath.relative_to(base_folder)
            folder_parts = [p.lower() for p in rel_path.parts[:-1]]
        except ValueError:
            folder_parts = []

        doc_info = DocumentInfo(
            path=filepath,
            filename=filename,
            suffix=suffix,
            category=DocumentCategory.EXTENDED,
            doc_type="unknown",
            size_bytes=filepath.stat().st_size if filepath.exists() else 0,
        )
        doc_info.date_hint = self._extract_date(filename)

        for skip_folder in self.folder_hints.get("skip", []):
            if any(skip_folder in part for part in folder_parts):
                doc_info.category = DocumentCategory.SKIP
                return doc_info

        doc_type = self._detect_doc_type(filename, suffix)
        doc_info.doc_type = doc_type

        is_root = len(folder_parts) == 0
        if is_root and not filename.startswith("_"):
            doc_info.category = DocumentCategory.CORE
            if doc_type == "unknown":
                doc_info.doc_type = "log" if suffix in [".docx", ".doc"] else "root_document"
            return doc_info

        if doc_type in ["log", "intake_sheet"]:
            doc_info.category = (
                DocumentCategory.CORE
                if self._is_in_report_period(doc_info.date_hint)
                else DocumentCategory.EXTENDED
            )
        elif doc_type == "approval":
            doc_info.category = DocumentCategory.CORE
        elif doc_type == "domain_report":
            doc_info.category = DocumentCategory.CORE
        elif doc_type in ["professional_report", "institutional_report"]:
            doc_info.category = (
                DocumentCategory.TIER2
                if self._is_recent(doc_info.date_hint, years=2)
                else DocumentCategory.EXTENDED
            )
        elif doc_type == "mail":
            doc_info.category = (
                DocumentCategory.TIER2
                if self._is_in_report_period(doc_info.date_hint)
                else DocumentCategory.EXTENDED
            )

        for part in folder_parts:
            if any(hint in part for hint in self.folder_hints.get("core", [])):
                if doc_info.category != DocumentCategory.SKIP:
                    if self._is_in_report_period(doc_info.date_hint):
                        doc_info.category = DocumentCategory.CORE
            elif any(hint in part for hint in self.folder_hints.get("extended", [])):
                if doc_info.category not in [DocumentCategory.SKIP, DocumentCategory.CORE]:
                    doc_info.category = DocumentCategory.EXTENDED

        return doc_info

    def _detect_doc_type(self, filename: str, suffix: str) -> str:
        """Erkennt den Dokumenttyp anhand des Dateinamens."""
        if suffix in [".msg", ".eml"]:
            return "mail"
        for doc_type, patterns in self._compiled_patterns.items():
            for pattern in patterns:
                if pattern.search(filename):
                    return doc_type
        return "unknown"

    def _extract_date(self, filename: str) -> Optional[datetime]:
        """Versucht ein Datum aus dem Dateinamen zu extrahieren."""
        patterns = [
            r"\b(\d{4})-(\d{2})-(\d{2})\b",
            r"\b(\d{2})\.(\d{2})\.(\d{4})\b",
            r"\b(\d{4})-(\d{2})\b",
            r"\b(\d{2})-(\d{4})\b",
            r"_(\d{4})(?:[_\.\s]|$)",
        ]
        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                groups = match.groups()
                try:
                    if len(groups) == 3:
                        if len(groups[0]) == 4:
                            return datetime(int(groups[0]), int(groups[1]), int(groups[2]))
                        return datetime(int(groups[2]), int(groups[1]), int(groups[0]))
                    if len(groups) == 2:
                        if len(groups[0]) == 4:
                            return datetime(int(groups[0]), int(groups[1]), 1)
                        return datetime(int(groups[1]), int(groups[0]), 1)
                    if len(groups) == 1:
                        return datetime(int(groups[0]), 6, 15)
                except (ValueError, IndexError):
                    pass
        return None

    def _is_in_report_period(self, date: Optional[datetime]) -> bool:
        if date is None:
            return True
        return date >= self.period_start

    def _is_recent(self, date: Optional[datetime], years: int = 2) -> bool:
        if date is None:
            return True
        cutoff = self.reference_date - timedelta(days=years * 365)
        return date >= cutoff

    # ─────────────────────────────────────────────────────────────
    # Interface 2: Text extrahieren und Bundle bauen
    # ─────────────────────────────────────────────────────────────

    def extract_bundle(
        self,
        documents: List[DocumentInfo],
        include_extended: bool = False,
        anonym_profile: Optional[object] = None,
    ) -> TextBundle:
        """Extrahiert Text aus den kategorisierten Dokumenten."""
        bundle = TextBundle()

        sorted_replacements: List[Tuple[str, str]] = []
        if anonym_profile is not None and hasattr(anonym_profile, "mappings"):
            all_replacements: Dict[str, str] = {}
            for category in anonym_profile.mappings.values():
                all_replacements.update(category)
            sorted_replacements = sorted(
                all_replacements.items(), key=lambda x: len(x[0]), reverse=True
            )

        def anonymize_text(text: str) -> str:
            if not sorted_replacements:
                return text
            for old, new in sorted_replacements:
                pattern = re.compile(r"\b" + re.escape(old) + r"\b")
                text = pattern.sub(lambda m: new, text)
            return text

        high_prio_parts, medium_prio_parts, low_prio_parts = [], [], []
        tier2_parts, extended_parts = [], []
        ten_years_ago = datetime.now() - timedelta(days=10 * 365)

        for doc in documents:
            if doc.category == DocumentCategory.SKIP:
                continue
            if doc.category == DocumentCategory.EXTENDED and not include_extended:
                continue

            text = self._extract_text_from_file(doc.path)
            if not text or text.startswith("[FEHLER"):
                continue
            text = anonymize_text(text)

            doc.text_extracted = True
            doc.text_length = len(text)
            bundle.files_processed += 1

            anon_filename = anonymize_text(doc.filename)
            content = f"--- {doc.doc_type.upper()}: {anon_filename} ---\n{text}"

            if doc.category == DocumentCategory.CORE:
                if doc.doc_type in ("log", "intake_sheet", "approval"):
                    high_prio_parts.append(content)
                elif doc.doc_type == "professional_report" and doc.date_hint and doc.date_hint < ten_years_ago:
                    low_prio_parts.append(content)
                else:
                    medium_prio_parts.append(content)
            elif doc.category == DocumentCategory.TIER2:
                if doc.doc_type == "professional_report" and doc.date_hint and doc.date_hint < ten_years_ago:
                    low_prio_parts.append(content)
                else:
                    tier2_parts.append(content)
            else:
                extended_parts.append(content)

        core_sections = []
        if high_prio_parts:
            core_sections.append(
                "=== HOHE PRIORITAET: Aktuelle Protokolle, Dokumentation, "
                "Genehmigungen, Aufnahmeblatt ==="
            )
            core_sections.extend(high_prio_parts)
        if medium_prio_parts:
            core_sections.append(
                "\n=== MITTLERE PRIORITAET: Fachberichte ==="
            )
            core_sections.extend(medium_prio_parts)
        if low_prio_parts:
            core_sections.append("\n=== NIEDRIGE PRIORITAET: Aeltere Berichte (>10 Jahre) ===")
            core_sections.extend(low_prio_parts)

        bundle.core_text = "\n\n".join(core_sections)
        bundle.tier2_text = "\n\n".join(tier2_parts)
        bundle.extended_text = "\n\n".join(extended_parts)
        bundle.total_length = len(bundle.core_text) + len(bundle.tier2_text)
        return bundle

    def save_bundle(self, bundle: TextBundle, output_path: str, combined: bool = True) -> Path:
        """Speichert ein Bundle als Textdatei (atomar, ueberschreibt nie)."""
        path = Path(output_path).expanduser().resolve()
        path.parent.mkdir(parents=True, exist_ok=True)

        def publish_new(target: Path, content: str) -> None:
            descriptor, temp_name = tempfile.mkstemp(prefix=f".{target.name}.", dir=target.parent)
            temporary = Path(temp_name)
            try:
                with os.fdopen(descriptor, "w", encoding="utf-8", newline="\n") as handle:
                    handle.write(content)
                    handle.flush()
                    os.fsync(handle.fileno())
                os.link(temporary, target)
            finally:
                temporary.unlink(missing_ok=True)

        if combined:
            content = "=== CORE DOKUMENTE ===\n\n" + bundle.core_text
            content += "\n\n=== TIER 2 DOKUMENTE ===\n\n" + bundle.tier2_text
            publish_new(path, content)
        else:
            core_path = path.with_suffix(".core.txt")
            tier2_path = path.with_suffix(".tier2.txt")
            if core_path.exists() or tier2_path.exists():
                raise FileExistsError("Bundle-Zieldatei existiert bereits")
            publish_new(core_path, bundle.core_text)
            try:
                publish_new(tier2_path, bundle.tier2_text)
            except Exception:
                core_path.unlink(missing_ok=True)
                raise
        return path

    # ─────────────────────────────────────────────────────────────
    # Interface 3: Bundle-Verzeichnis lesen
    # ─────────────────────────────────────────────────────────────

    def read_bundle_dir(
        self, bundle_path: Path, filename_filter: Optional[str] = None, include_extended: bool = False
    ) -> str:
        """Liest Inhalte aus einem bereits (extern) anonymisierten Bundle-Verzeichnis."""
        bundle_path = bundle_path.expanduser().resolve()
        if not bundle_path.is_dir() or bundle_path.is_symlink():
            raise ValueError("Bundle-Verzeichnis ist ungültig")
        folders = [bundle_path / "core"]
        if include_extended and (bundle_path / "extended").exists():
            folders.append(bundle_path / "extended")

        all_content = []
        total_bytes = 0
        accepted_files = 0

        for folder in folders:
            if not folder.exists():
                continue
            for filepath in sorted(folder.rglob("*")):
                if filepath.is_symlink() or not filepath.is_file():
                    continue
                if filepath.suffix.lower() not in BUNDLE_READ_EXTENSIONS:
                    continue
                if filepath.name.startswith(".") or filepath.name.startswith("_"):
                    continue
                if "output" in filepath.parts:
                    continue

                resolved = filepath.resolve()
                try:
                    resolved.relative_to(bundle_path)
                except ValueError as exc:
                    raise ValueError("Bundle-Datei verlässt den Bundle-Root") from exc
                size = resolved.stat().st_size
                if size > MAX_DOCUMENT_BYTES:
                    raise ValueError("Bundle-Datei überschreitet das Größenlimit")
                total_bytes += size
                accepted_files += 1
                if accepted_files > MAX_DOCUMENT_FILES or total_bytes > MAX_TOTAL_DOCUMENT_BYTES:
                    raise ValueError("Bundle überschreitet das Ressourcenlimit")

                if filename_filter and filename_filter.lower() not in filepath.name.lower():
                    continue

                text = self._extract_text_from_file(filepath)
                if text.strip():
                    rel_path = filepath.relative_to(bundle_path)
                    all_content.append(f"--- Quelle: {rel_path} ---\n{text}")

        if not all_content:
            return "[KEINE DATEN] Keine passenden Dokumente im Bundle gefunden."
        return "\n\n".join(all_content)

    def list_bundle_files(self, bundle_path: Path) -> List[dict]:
        """Listet alle verfuegbaren Dateien in einem Bundle auf."""
        if not bundle_path.exists():
            return []
        files = []
        for filepath in sorted(bundle_path.rglob("*")):
            if not filepath.is_file():
                continue
            if filepath.name.startswith("."):
                continue
            if "output" in filepath.parts:
                continue
            rel_path = filepath.relative_to(bundle_path)
            files.append({
                "name": filepath.name,
                "path": str(rel_path),
                "size": filepath.stat().st_size,
                "type": filepath.suffix.lower(),
            })
        return files

    # ─────────────────────────────────────────────────────────────
    # Text-Extraktion
    # ─────────────────────────────────────────────────────────────

    def _extract_text_from_file(self, filepath: Path) -> str:
        suffix = filepath.suffix.lower()
        filepath_str = str(filepath)
        try:
            if suffix == ".docx":
                return self._extract_docx(filepath_str)
            elif suffix == ".doc":
                return self._extract_doc(filepath_str)
            elif suffix in [".txt", ".md"]:
                return self._extract_txt(filepath_str)
            elif suffix == ".pdf":
                return self._extract_pdf(filepath_str)
            elif suffix in [".xlsx", ".xls"]:
                return self._extract_excel(filepath_str)
            elif suffix == ".msg":
                return self._extract_msg(filepath_str)
            elif suffix == ".eml":
                return self._extract_eml(filepath_str)
            else:
                return f"[Nicht unterstuetztes Format: {suffix}]"
        except Exception as e:
            return f"[FEHLER bei {filepath.name}: {e}]"

    def _extract_docx(self, filepath: str) -> str:
        try:
            from docx import Document

            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return "[python-docx nicht installiert]"
        except Exception as e:
            return f"[DOCX-Fehler: {e}]"

    def _extract_doc(self, filepath: str) -> str:
        """Extrahiert Text aus altem Word-Format (.doc) via antiword oder LibreOffice."""
        import subprocess
        import shutil

        if shutil.which("antiword"):
            try:
                result = subprocess.run(["antiword", filepath], capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except Exception:
                pass

        if shutil.which("soffice"):
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    result = subprocess.run(
                        ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, filepath],
                        capture_output=True, timeout=60,
                    )
                    if result.returncode == 0:
                        txt_file = Path(tmpdir) / (Path(filepath).stem + ".txt")
                        if txt_file.exists():
                            return txt_file.read_text(encoding="utf-8", errors="replace")
            except Exception:
                pass

        try:
            import textract

            text = textract.process(filepath).decode("utf-8", errors="replace")
            return text.strip()
        except ImportError:
            pass
        except Exception:
            pass

        return "[.doc-Extraktion fehlgeschlagen - antiword/LibreOffice/textract nicht verfuegbar]"

    def _extract_txt(self, filepath: str) -> str:
        path = Path(filepath)
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    def _extract_pdf(self, filepath: str) -> str:
        text_parts = []
        needs_ocr = False
        try:
            from pypdf import PdfReader

            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    text_parts.append(page_text)
                else:
                    needs_ocr = True
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        try:
            import pdfplumber

            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        needs_ocr = True
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception:
            pass

        try:
            import fitz

            doc = fitz.open(filepath)
            fitz_parts = []
            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    fitz_parts.append(page_text)
            doc.close()
            if fitz_parts:
                return "\n".join(fitz_parts)
        except ImportError:
            pass
        except Exception:
            pass

        if needs_ocr:
            return self._extract_pdf_ocr(filepath)
        return "[PDF-Extraktion fehlgeschlagen: pypdf, pdfplumber und PyMuPDF nicht verfuegbar]"

    def _extract_pdf_ocr(self, filepath: str) -> str:
        """OCR fuer Bild-PDFs -- erfordert einen extern bereitgestellten Adapter
        (Modulname/Schnittstelle liegt beim Overlay, hier bewusst kein
        hartcodiertes Fremdmodul)."""
        try:
            from c_ocr_engine import is_ocr_available, ocr_pdf

            if not is_ocr_available():
                return "[OCR nicht verfuegbar - Tesseract nicht installiert]"
            text = ocr_pdf(filepath)
            return text if text else "[OCR: Kein Text erkannt]"
        except ImportError:
            return "[OCR-Adapter nicht installiert]"
        except Exception as exc:
            return f"[OCR-Fehler ({type(exc).__name__})]"

    def _extract_excel(self, filepath: str) -> str:
        try:
            import openpyxl

            wb = openpyxl.load_workbook(filepath, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"[Tabelle: {sheet_name}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)
        except ImportError:
            return "[openpyxl nicht installiert]"

    def _extract_msg(self, filepath: str) -> str:
        try:
            import extract_msg

            msg = extract_msg.Message(filepath)
            parts = []
            if msg.date:
                parts.append(f"Datum: {msg.date}")
            if msg.sender:
                parts.append(f"Von: {msg.sender}")
            if msg.to:
                parts.append(f"An: {msg.to}")
            if msg.subject:
                parts.append(f"Betreff: {msg.subject}")
            parts.append("")
            if msg.body:
                parts.append(msg.body)

            if hasattr(msg, "attachments") and msg.attachments:
                for attach in msg.attachments:
                    attach_name = "unbekannt"
                    try:
                        attach_name = attach.longFilename or attach.shortFilename or "unbekannt"
                        suffix = Path(attach_name).suffix.lower()
                        if suffix == ".pdf":
                            tmp_path = None
                            try:
                                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                                    tmp.write(attach.data)
                                    tmp_path = tmp.name
                                parts.append(f"\n[Anhang: {attach_name}]")
                                parts.append(self._extract_pdf(tmp_path))
                            finally:
                                if tmp_path:
                                    Path(tmp_path).unlink(missing_ok=True)
                        elif suffix == ".docx":
                            tmp_path = None
                            try:
                                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                                    tmp.write(attach.data)
                                    tmp_path = tmp.name
                                parts.append(f"\n[Anhang: {attach_name}]")
                                parts.append(self._extract_docx(tmp_path))
                            finally:
                                if tmp_path:
                                    Path(tmp_path).unlink(missing_ok=True)
                        elif suffix in [".txt", ".md"]:
                            parts.append(f"\n[Anhang: {attach_name}]")
                            try:
                                text = attach.data.decode("utf-8")
                            except Exception:
                                text = attach.data.decode("latin-1")
                            parts.append(text)
                    except Exception as e:
                        parts.append(f"\n[Anhang {attach_name}: Fehler - {e}]")

            msg.close()
            return "\n".join(parts)
        except ImportError:
            return "[extract-msg nicht installiert]"

    def _extract_eml(self, filepath: str) -> str:
        from email import policy
        from email.parser import BytesParser

        with open(filepath, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []
        if msg["Date"]:
            parts.append(f"Datum: {msg['Date']}")
        if msg["From"]:
            parts.append(f"Von: {msg['From']}")
        if msg["To"]:
            parts.append(f"An: {msg['To']}")
        if msg["Subject"]:
            parts.append(f"Betreff: {msg['Subject']}")
        parts.append("")

        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    charset = part.get_content_charset() or "utf-8"
                    try:
                        body = part.get_payload(decode=True).decode(charset)
                        parts.append(body)
                        break
                    except Exception:
                        pass
        else:
            charset = msg.get_content_charset() or "utf-8"
            try:
                body = msg.get_payload(decode=True).decode(charset)
                parts.append(body)
            except Exception:
                parts.append(str(msg.get_payload()))

        return "\n".join(parts)


# ═══════════════════════════════════════════════════════════════
# Convenience-Funktion
# ═══════════════════════════════════════════════════════════════


def scan_and_extract(
    folder_path: str,
    output_path: Optional[str] = None,
    include_extended: bool = False,
    *,
    allow_unanonymized: bool = False,
) -> Tuple[TextBundle, None]:
    """
    Explizite lokale Klartextverarbeitung (KEINE Anonymisierung).

    Diese Funktion erzeugt KEIN Anonymisierungsprofil -- fuer LLM-Nutzung
    ist stattdessen ReportWorkflow.prepare() (mode="anonymized") zu
    verwenden.

    Args:
        folder_path: Pfad zum Quellordner
        output_path: Optional - Pfad zum Speichern des Bundles
        include_extended: EXTENDED-Dokumente einbeziehen
        allow_unanonymized: Muss explizit True sein (Fail-closed-Guard)

    Returns:
        (TextBundle, None)
    """
    if not allow_unanonymized:
        raise PermissionError(
            "Klartext-Extraktion erfordert allow_unanonymized=True; "
            "fuer LLM-Nutzung ReportWorkflow.prepare(mode='anonymized') verwenden."
        )

    pipeline = DocumentPipeline()
    result = pipeline.scan_folder(folder_path)
    if result.errors:
        raise RuntimeError("Dokument-Scan abgebrochen: " + "; ".join(result.errors))
    bundle = pipeline.extract_bundle(result.documents, include_extended=include_extended, anonym_profile=None)

    if output_path:
        pipeline.save_bundle(bundle, output_path)

    return bundle, None
