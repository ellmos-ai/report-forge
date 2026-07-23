#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
Copyright (c) 2026 Lukas Geiger

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.

render_utils.py -- Domaenen-neutrale Render-Hilfsfunktionen
=============================================================

Portiert aus dem privaten foerderplaner-Skill (generator/generator.py,
"Renderteil"): Funktionen, die beim Befuellen einer Word-Vorlage aus
JSON-Daten in JEDEM Berichtsdomaenen-Kontext gebraucht werden, ohne
irgendeine domaenenspezifische Feld- oder Vokabularannahme.

Version: 1.0.0
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, List


# ===================================================================
# Ergebnis-Datenklasse
# ===================================================================


@dataclass
class RenderResult:
    """Ergebnis eines Render-/Fuellvorgangs (Phase 3: Word-Vorlage befuellen)."""

    success: bool = False
    output_path: str = ""
    template_filled: bool = False
    json_generated: bool = False
    source_text_length: int = 0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# ===================================================================
# XML-Platzhalter-Ersetzung ueber Run-/Element-Grenzen hinweg
# ===================================================================


def replace_complex_placeholder_in_xml(doc: Any, search_text: str, replacement: str) -> bool:
    """
    Ersetzt Text, der ueber mehrere Word-XML-Text-Elemente (w:t) verteilt ist.

    Word teilt sichtbaren Text haeufig in mehrere interne Runs/Elemente auf
    (Autokorrektur, Formatierungswechsel, Bearbeitungshistorie). Eine
    Ersetzung, die nur ein einzelnes Element betrachtet, findet solche
    Platzhalter nicht -- dieses Verfahren sammelt den zusammenhaengenden
    Text ueber alle w:t-Elemente eines Paragraphen/einer Zelle und schreibt
    die Ersetzung an der richtigen Stelle zurueck.

    Args:
        doc: Ein python-docx Document
        search_text: Gesuchter Text (kann ueber mehrere XML-Elemente verteilt sein)
        replacement: Ersetzungstext

    Returns:
        True wenn mindestens eine Ersetzung stattfand.
    """
    from docx.oxml.ns import qn

    replaced = False

    def _process(t_elements):
        nonlocal replaced
        full_text = "".join(t.text or "" for t in t_elements)
        if search_text not in full_text:
            return
        start_idx = full_text.find(search_text)
        end_idx = start_idx + len(search_text)

        char_idx = 0
        first_t = None
        for t in t_elements:
            t_text = t.text or ""
            t_start = char_idx
            t_end = char_idx + len(t_text)

            if t_start <= start_idx < t_end and first_t is None:
                first_t = t
                prefix = t_text[: start_idx - t_start]
                suffix_start = min(end_idx - t_start, len(t_text))
                suffix = t_text[suffix_start:] if suffix_start < len(t_text) else ""
                t.text = prefix + replacement + suffix
                replaced = True
            elif first_t is not None and t_start < end_idx:
                if t_end <= end_idx:
                    t.text = ""
                else:
                    t.text = t_text[end_idx - t_start :]

            char_idx = t_end

    for para in doc.paragraphs:
        _process(list(para._element.iter(qn("w:t"))))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process(list(para._element.iter(qn("w:t"))))

    return replaced


# ===================================================================
# Umlaut-Normalisierung (ASCII-Transliteration <-> echte Umlaute)
# ===================================================================

# Generische Ausnahmewoerter, in denen "ae"/"oe"/"ue" KEIN Umlaut-Digraph
# ist (z.B. "aktuell", "Steuer", "Museum"). Domaenen koennen ueber
# extra_exceptions weitere, eigene Ausnahmen ergaenzen (z.B. Eigennamen
# oder Fachbegriffe ihres Vokabulars), ohne diese Liste hier zu aendern.
_DEFAULT_NO_REPLACE_PATTERNS = {
    "aktuell", "abenteuer", "steuer", "feuer", "teuer", "heuer", "ungeheuer",
    "euer", "neue", "treue", "reue", "freue", "bauer", "mauer", "lauer",
    "sauer", "trauer", "dauer", "genauer", "blauer", "schauer", "grauer",
    "rauer", "museum", "duell", "fuel", "cruel", "samuel", "manuel",
    "virtuell", "eventuell", "individuell", "sexuell", "visuell", "rituell",
    "intellektuell", "spirituell", "residuell", "queue", "sequel", "tissue",
    "venue", "revenue", "israel", "michael", "raphael", "aerob", "aero",
    "poem", "poet", "phoenix", "aloe", "shoe", "joe", "noel",
}


def fix_umlauts_in_values(data: Any, extra_exceptions: frozenset[str] = frozenset()) -> Any:
    """
    Konvertiert ASCII-Umlaut-Transliterationen (ae/oe/ue) in echte deutsche
    Umlaute (ä/ö/ü) in allen String-WERTEN einer beliebig verschachtelten
    JSON-artigen Struktur. Keys bleiben unveraendert (siehe
    normalize_umlaut_keys fuer die Gegenrichtung).

    Wortbasierte Pruefung vermeidet False Positives ("aktuell" bleibt
    "aktuell", nicht "aktäll"). extra_exceptions erlaubt domaenenspezifische
    Zusatz-Ausnahmen (z.B. Eigennamen/Fachbegriffe), ohne die generische
    Kernliste zu veraendern.
    """
    exceptions = _DEFAULT_NO_REPLACE_PATTERNS | {e.lower() for e in extra_exceptions}

    def _should_replace(text_lower, pos, digraph_len):
        digraph = text_lower[pos : pos + digraph_len]
        if digraph == "ue" and text_lower[max(0, pos - 2) : pos] == "ng":
            return False
        start = pos
        while start > 0 and text_lower[start - 1].isalpha():
            start -= 1
        end = pos + digraph_len
        while end < len(text_lower) and text_lower[end].isalpha():
            end += 1
        word = text_lower[start:end]
        return not any(exc in word for exc in exceptions)

    def fix_text(text):
        if not isinstance(text, str):
            return text
        result = list(text)
        text_lower = text.lower()
        replacements = [
            ("ae", "ä"), ("Ae", "Ä"),
            ("oe", "ö"), ("Oe", "Ö"),
            ("ue", "ü"), ("Ue", "Ü"),
        ]
        for ascii_form, umlaut in replacements:
            offset = 0
            while True:
                pos = text_lower.find(ascii_form.lower(), offset)
                if pos == -1:
                    break
                if _should_replace(text_lower, pos, len(ascii_form)):
                    if text[pos].isupper():
                        replacement = umlaut.upper() if len(umlaut) == 1 else umlaut
                    else:
                        replacement = umlaut.lower() if len(umlaut) == 1 else umlaut
                    result[pos : pos + len(ascii_form)] = list(replacement)
                    text = "".join(result)
                    text_lower = text.lower()
                    offset = pos + len(replacement)
                else:
                    offset = pos + len(ascii_form)
        return "".join(result)

    if isinstance(data, dict):
        return {k: fix_umlauts_in_values(v, extra_exceptions) for k, v in data.items()}
    if isinstance(data, list):
        return [fix_umlauts_in_values(item, extra_exceptions) for item in data]
    if isinstance(data, str):
        return fix_text(data)
    return data


def normalize_umlaut_keys(data: Any) -> Any:
    """Kanonisiert Objekt-KEYS auf ASCII (ä->ae usw.), ohne Werte zu veraendern."""
    umlaut_map = str.maketrans(
        {
            "ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss",
            "Ä": "Ae", "Ö": "Oe", "Ü": "Ue",
        }
    )
    if isinstance(data, dict):
        return {
            str(key).translate(umlaut_map): normalize_umlaut_keys(value)
            for key, value in data.items()
        }
    if isinstance(data, list):
        return [normalize_umlaut_keys(item) for item in data]
    return data
